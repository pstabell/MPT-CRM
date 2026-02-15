"""
Contract Generation Service for MPT-CRM
=======================================

Service for generating contracts by merging:
- NDA (pages 1-2)
- SOW with auto-filled client info (pages 3-4) 
- Proposal as Exhibit A (remaining pages)

Dependencies: python-docx, PyMuPDF (fitz), reportlab
"""

import os
import fitz  # PyMuPDF
from datetime import datetime
import uuid
import tempfile
from pathlib import Path
import logging
from typing import Dict, Optional, List
from io import BytesIO

# Import docx for Word template processing
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Install with: pip install python-docx")

# Import reportlab for PDF creation from docx
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("reportlab already listed in requirements.txt")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Template paths
CONTRACTS_DIR = r"C:\Users\Patri\Metro Point Technology\Metro Point Technology - Documents\SALES\CONTRACTS"
NDA_TEMPLATE_PATH = os.path.join(CONTRACTS_DIR, "Mutual Non-Disclosure Agreement (NDA).docx")
SOW_TEMPLATE_PATH = os.path.join(CONTRACTS_DIR, "Statement of Work (SOW) TEMPLATE.docx")

class ContractGenerationError(Exception):
    """Custom exception for contract generation errors"""
    pass

def verify_dependencies():
    """Verify required dependencies are available"""
    missing = []
    if not DOCX_AVAILABLE:
        missing.append("python-docx")
    if not REPORTLAB_AVAILABLE:
        missing.append("reportlab") 
    
    if missing:
        raise ContractGenerationError(f"Missing dependencies: {', '.join(missing)}. Install with: pip install {' '.join(missing)}")
    
    return True

def verify_templates():
    """Verify contract templates exist"""
    if not os.path.exists(NDA_TEMPLATE_PATH):
        raise ContractGenerationError(f"NDA template not found: {NDA_TEMPLATE_PATH}")
    
    if not os.path.exists(SOW_TEMPLATE_PATH):
        raise ContractGenerationError(f"SOW template not found: {SOW_TEMPLATE_PATH}")
    
    logger.info("✅ Contract templates verified")
    return True

def auto_fill_sow_template(project_data: Dict) -> str:
    """
    Auto-fill SOW template with client data
    
    Args:
        project_data: Dictionary containing project and client information
        
    Returns:
        Path to the filled SOW document
        
    Expected project_data structure:
    {
        'client_name': str,
        'project_name': str,  
        'budget': float,
        'contact_name': str,
        'contact_email': str
    }
    """
    verify_dependencies()
    verify_templates()
    
    logger.info("🔧 Auto-filling SOW template...")
    
    try:
        # Load SOW template
        doc = Document(SOW_TEMPLATE_PATH)
        
        # Auto-fill fields
        current_date = datetime.now().strftime("%B %d, %Y")
        
        # Replacement mappings
        replacements = {
            '[Date]': current_date,
            '[Client Name]': project_data.get('client_name', '[Client Name]'),
            '[Insert Project Title]': project_data.get('project_name', '[Project Title]'),
            '$[Amount]': f"${project_data.get('budget', 0):,.2f}",
            '[Contact Name]': project_data.get('contact_name', '[Contact Name]'),
            '[Contact Email]': project_data.get('contact_email', '[Contact Email]')
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    logger.info(f"   Replaced '{old_text}' with '{new_text}'")
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_text, new_text in replacements.items():
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)
                            logger.info(f"   Replaced '{old_text}' with '{new_text}' in table")
        
        # Save filled template
        temp_dir = Path("temp_documents")
        temp_dir.mkdir(exist_ok=True)
        
        filled_sow_path = temp_dir / f"SOW_filled_{uuid.uuid4().hex[:8]}.docx"
        doc.save(str(filled_sow_path))
        
        logger.info(f"✅ SOW auto-filled and saved to: {filled_sow_path}")
        return str(filled_sow_path)
        
    except Exception as e:
        logger.error(f"❌ Error auto-filling SOW: {e}")
        raise ContractGenerationError(f"Failed to auto-fill SOW template: {e}")

def convert_docx_to_pdf(docx_path: str) -> str:
    """
    Convert DOCX file to PDF
    
    Note: This is a simplified conversion. For production use, consider:
    - python-docx2pdf (Windows only, requires Word)
    - libreoffice --convert-to pdf (requires LibreOffice)
    - pandoc (requires pandoc installation)
    
    This implementation creates a basic PDF from docx content.
    """
    logger.info(f"📄 Converting DOCX to PDF: {docx_path}")
    
    try:
        # Load docx
        doc = Document(docx_path)
        
        # Create PDF path
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        # Create PDF using reportlab
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.utils import simpleSplit
        
        c = canvas.Canvas(pdf_path, pagesize=A4)
        width, height = A4
        
        y_position = height - 50  # Start near top
        line_height = 14
        margin = 50
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Split long lines to fit page width
                lines = simpleSplit(paragraph.text, 'Helvetica', 10, width - 2*margin)
                
                for line in lines:
                    if y_position < 50:  # Start new page
                        c.showPage()
                        y_position = height - 50
                    
                    c.drawString(margin, y_position, line)
                    y_position -= line_height
                
                # Add space after paragraph
                y_position -= 5
        
        c.save()
        
        logger.info(f"✅ PDF created: {pdf_path}")
        return pdf_path
        
    except Exception as e:
        logger.error(f"❌ Error converting DOCX to PDF: {e}")
        # For now, return the docx path and let the caller handle it
        raise ContractGenerationError(f"Failed to convert DOCX to PDF: {e}")

def get_proposal_pdf_path(project_data: Dict) -> Optional[str]:
    """
    Get the path to the proposal PDF for this project
    
    This should integrate with the existing proposal system to find 
    the associated proposal PDF file.
    """
    # TODO: Implement integration with proposal system
    # For now, return None if no proposal found
    
    project_id = project_data.get('id')
    deal_id = project_data.get('deal_id')
    
    logger.info(f"🔍 Looking for proposal PDF for project {project_id}, deal {deal_id}")
    
    # Check if there's a folder_url or proposal path in project data
    folder_url = project_data.get('folder_url')
    if folder_url:
        logger.info(f"   Found SharePoint folder: {folder_url}")
        # TODO: Download proposal from SharePoint if needed
    
    # For development/testing, return None
    logger.warning("⚠️  Proposal PDF integration not implemented yet")
    return None

def merge_contract_pdfs(nda_pdf_path: str, sow_pdf_path: str, proposal_pdf_path: Optional[str] = None) -> str:
    """
    Merge NDA + SOW + Proposal PDFs into a single contract document
    
    Args:
        nda_pdf_path: Path to NDA PDF (pages 1-2)
        sow_pdf_path: Path to SOW PDF (pages 3-4)
        proposal_pdf_path: Optional path to proposal PDF (Exhibit A)
        
    Returns:
        Path to merged contract PDF
    """
    logger.info("🔄 Merging contract PDFs...")
    
    try:
        # Create output path
        temp_dir = Path("temp_documents")
        temp_dir.mkdir(exist_ok=True)
        
        contract_filename = f"Contract_Merged_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        merged_pdf_path = temp_dir / contract_filename
        
        # Create new PDF document
        merged_doc = fitz.open()
        
        # Add NDA (pages 1-2)
        if os.path.exists(nda_pdf_path):
            nda_doc = fitz.open(nda_pdf_path)
            merged_doc.insert_pdf(nda_doc)
            nda_doc.close()
            logger.info("   ✅ Added NDA pages")
        else:
            logger.warning(f"   ⚠️  NDA PDF not found: {nda_pdf_path}")
        
        # Add SOW (pages 3-4)  
        if os.path.exists(sow_pdf_path):
            sow_doc = fitz.open(sow_pdf_path)
            merged_doc.insert_pdf(sow_doc)
            sow_doc.close()
            logger.info("   ✅ Added SOW pages")
        else:
            logger.warning(f"   ⚠️  SOW PDF not found: {sow_pdf_path}")
        
        # Add Proposal as Exhibit A (remaining pages)
        if proposal_pdf_path and os.path.exists(proposal_pdf_path):
            proposal_doc = fitz.open(proposal_pdf_path)
            merged_doc.insert_pdf(proposal_doc)
            proposal_doc.close()
            logger.info("   ✅ Added Proposal as Exhibit A")
        else:
            logger.info("   ℹ️  No proposal PDF to add")
        
        # Save merged document
        merged_doc.save(str(merged_pdf_path))
        merged_doc.close()
        
        # Get page count
        with fitz.open(str(merged_pdf_path)) as doc:
            page_count = len(doc)
        
        logger.info(f"✅ Contract merged successfully: {merged_pdf_path} ({page_count} pages)")
        return str(merged_pdf_path)
        
    except Exception as e:
        logger.error(f"❌ Error merging PDFs: {e}")
        raise ContractGenerationError(f"Failed to merge contract PDFs: {e}")

def generate_contract(project_data: Dict) -> Dict:
    """
    Main function to generate a complete contract package
    
    Args:
        project_data: Project and client information
        
    Returns:
        Dictionary with contract generation results:
        {
            'success': bool,
            'contract_pdf_path': str,
            'page_count': int,
            'components': {
                'nda_pdf': str,
                'sow_pdf': str, 
                'proposal_pdf': str or None
            },
            'error': str or None
        }
    """
    logger.info("🚀 Starting contract generation...")
    logger.info(f"   Client: {project_data.get('client_name')}")
    logger.info(f"   Project: {project_data.get('project_name')}")
    
    try:
        # Step 1: Verify everything is ready
        verify_dependencies()
        verify_templates()
        
        # Step 2: Auto-fill SOW template
        filled_sow_docx = auto_fill_sow_template(project_data)
        
        # Step 3: Convert templates to PDF
        # For NDA, we need to convert the docx template to PDF
        nda_pdf_path = convert_docx_to_pdf(NDA_TEMPLATE_PATH)
        sow_pdf_path = convert_docx_to_pdf(filled_sow_docx)
        
        # Step 4: Get proposal PDF (if available)
        proposal_pdf_path = get_proposal_pdf_path(project_data)
        
        # Step 5: Merge all PDFs
        contract_pdf_path = merge_contract_pdfs(nda_pdf_path, sow_pdf_path, proposal_pdf_path)
        
        # Step 6: Get final page count
        with fitz.open(contract_pdf_path) as doc:
            page_count = len(doc)
        
        result = {
            'success': True,
            'contract_pdf_path': contract_pdf_path,
            'page_count': page_count,
            'components': {
                'nda_pdf': nda_pdf_path,
                'sow_pdf': sow_pdf_path,
                'proposal_pdf': proposal_pdf_path
            },
            'error': None
        }
        
        logger.info(f"🎉 Contract generation completed successfully!")
        logger.info(f"   Final contract: {contract_pdf_path} ({page_count} pages)")
        
        return result
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"❌ Contract generation failed: {error_msg}")
        
        return {
            'success': False,
            'contract_pdf_path': None,
            'page_count': 0,
            'components': {},
            'error': error_msg
        }

def cleanup_temp_files(file_paths: List[str]):
    """Clean up temporary files"""
    for file_path in file_paths:
        try:
            if file_path and os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"🗑️  Cleaned up: {file_path}")
        except Exception as e:
            logger.warning(f"⚠️  Could not clean up {file_path}: {e}")

# Development/Testing utilities
def test_contract_generation():
    """Test function to verify contract generation works"""
    test_project_data = {
        'id': 'test-project-123',
        'client_name': 'Acme Corporation',
        'project_name': 'Website Redesign Project',
        'budget': 15000.00,
        'contact_name': 'John Smith',
        'contact_email': 'john.smith@acme.com',
        'deal_id': 'test-deal-456'
    }
    
    logger.info("🧪 Running contract generation test...")
    result = generate_contract(test_project_data)
    
    if result['success']:
        logger.info("✅ Test completed successfully!")
        logger.info(f"   Contract: {result['contract_pdf_path']}")
        logger.info(f"   Pages: {result['page_count']}")
    else:
        logger.error(f"❌ Test failed: {result['error']}")
    
    return result

if __name__ == "__main__":
    # Run test when executed directly
    test_contract_generation()